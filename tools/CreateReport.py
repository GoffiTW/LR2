from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\danet\source\repos\Привязка и команды")
REPORT_PATH = ROOT / "ЛР_2_Данет_Данет.docx"
ASSETS = ROOT / "ReportAssets"


def set_font(run, size=None, bold=False):
    run.font.name = "Times New Roman"
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(document, text, style=None):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_font(run, 14)
    paragraph.paragraph_format.first_line_indent = Inches(0.35)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_font(run, 14)


def add_code(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def add_picture(document, file_name, caption):
    image_path = ASSETS / file_name
    if image_path.exists():
        document.add_picture(str(image_path), width=Inches(6.4))
        last = document.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    set_font(run, 12)
    run.italic = True


document = Document()
section = document.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Лабораторная работа №2\nПривязки и команды в WPF")
set_font(run, 18, True)

meta = document.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Отчет по выполненной работе")
set_font(run, 14)

document.add_page_break()

document.add_heading("Введение", level=1)
add_paragraph(
    document,
    "Цель работы: разработать WPF-приложение, демонстрирующее механизмы привязки данных, "
    "команды и триггеры. В приложении реализованы отдельные вкладки для привязки по умолчанию, "
    "двухсторонней, одноразовой и односторонней привязки, а также вкладка с вариантами триггеров.",
)
add_bullets(
    document,
    [
        "создать несколько вкладок и вынести содержимое каждой вкладки в отдельный UserControl;",
        "показать прямую привязку к элементам интерфейса и привязку через визуальную модель;",
        "реализовать кнопки через механизм ICommand;",
        "показать разницу UpdateSourceTrigger: PropertyChanged, LostFocus и Explicit;",
        "показать Trigger, DataTrigger, MultiDataTrigger и EventTrigger.",
    ],
)

document.add_heading("Разница в создании визуальных моделей в ветках", level=1)
add_paragraph(
    document,
    "В ветке a используется ручная реализация MVVM без вспомогательных библиотек и NuGet-пакетов. "
    "Для этого созданы классы ViewModelBase и RelayCommand. ViewModelBase реализует интерфейс "
    "INotifyPropertyChanged, а RelayCommand инкапсулирует логику ICommand.",
)
add_code(
    document,
    "public abstract class ViewModelBase : INotifyPropertyChanged\n"
    "{\n"
    "    public event PropertyChangedEventHandler PropertyChanged;\n"
    "    protected void OnPropertyChanged([CallerMemberName] string propertyName = null)\n"
    "        => PropertyChanged?.Invoke(this, new PropertyChangedEventArgs(propertyName));\n"
    "}",
)
add_paragraph(
    document,
    "В ветке b по заданию должен использоваться пакет CommunityToolkit.Mvvm. В этом подходе "
    "однотипный код INotifyPropertyChanged и ICommand заменяется атрибутами и базовыми классами "
    "из библиотеки. Это уменьшает объем шаблонного кода и делает визуальную модель компактнее.",
)
add_code(
    document,
    "public partial class MainViewModel : ObservableObject\n"
    "{\n"
    "    [ObservableProperty]\n"
    "    private string userName;\n\n"
    "    [RelayCommand]\n"
    "    private void SaveTwoWay() { }\n"
    "}",
)

document.add_heading("Виды привязок", level=1)

document.add_heading("Привязка по умолчанию", level=2)
add_paragraph(
    document,
    "Привязка по умолчанию зависит от свойства целевого элемента. Для TextBox.Text режим обычно "
    "двухсторонний с обновлением источника при потере фокуса. На вкладке показана прямая привязка "
    "через ElementName и привязка к свойствам MainViewModel.",
)
add_picture(document, "01_default_binding.png", "Рисунок 1 - Демонстрация привязки по умолчанию")
add_paragraph(document, "Реализация в ветке a: XAML использует обычный Binding и свойства MainViewModel.")
add_paragraph(document, "Реализация в ветке b: свойства объявляются через [ObservableProperty], привязка XAML остается такой же.")

document.add_heading("Двухсторонняя привязка", level=2)
add_paragraph(
    document,
    "TwoWay-связь синхронизирует источник и целевой элемент в обе стороны. Изменения в TextBox, "
    "Slider, CheckBox и ComboBox попадают в ViewModel, а изменения ViewModel сразу отображаются в интерфейсе.",
)
add_picture(document, "02_two_way_binding.png", "Рисунок 2 - Двухсторонняя привязка")
add_paragraph(document, "Реализация в ветке a: свойства UserName, Rating, NotificationsEnabled и SelectedTheme вызывают OnPropertyChanged.")
add_paragraph(document, "Реализация в ветке b: эти свойства создаются генератором CommunityToolkit.Mvvm.")

document.add_heading("Одноразовая привязка", level=2)
add_paragraph(
    document,
    "OneTime получает значение источника один раз при создании привязки. После изменения SessionCode "
    "командой одноразовый TextBlock сохраняет исходное значение, а OneWay-поле рядом обновляется.",
)
add_picture(document, "03_one_time_binding.png", "Рисунок 3 - Одноразовая привязка")
add_paragraph(document, "Реализация в ветке a: команда GenerateSessionCodeCommand изменяет свойство SessionCode.")
add_paragraph(document, "Реализация в ветке b: команда оформляется атрибутом [RelayCommand].")

document.add_heading("Односторонние привязки", level=2)
add_paragraph(
    document,
    "OneWay обновляет интерфейс при изменении источника, но пользовательский ввод не меняет источник. "
    "OneWayToSource работает наоборот: целевой элемент передает значение в источник, не читая его обратно.",
)
add_paragraph(
    document,
    "На этой же вкладке показана разница UpdateSourceTrigger. PropertyChanged обновляет источник сразу, "
    "LostFocus - после ухода фокуса, Explicit - только после выполнения команды обновления BindingExpression.",
)
add_picture(document, "04_one_way_bindings.png", "Рисунок 4 - Односторонние привязки и UpdateSourceTrigger")
add_paragraph(document, "Реализация в ветке a: команды IncreaseProgressCommand, DecreaseProgressCommand и ApplyExplicitBindingCommand описаны вручную.")
add_paragraph(document, "Реализация в ветке b: команды могут быть заменены методами с атрибутом [RelayCommand].")

document.add_heading("Триггеры", level=1)
add_paragraph(
    document,
    "Trigger реагирует на свойство самого элемента, DataTrigger - на данные из ViewModel, "
    "MultiDataTrigger - на несколько условий одновременно, EventTrigger - на событие и обычно "
    "запускает анимацию или storyboard.",
)
add_picture(document, "05_triggers.png", "Рисунок 5 - Варианты триггеров")
add_paragraph(
    document,
    "В приложении Property Trigger меняет блок при наведении мыши. DataTrigger меняет блок при включении "
    "опасного режима. MultiDataTrigger срабатывает только при опасном режиме и выбранном пункте 'Критический'. "
    "EventTrigger запускает анимацию прозрачности при MouseEnter.",
)

document.add_heading("Ссылка на проект в системе контроля версий", level=1)
add_paragraph(document, f"Локальный путь проекта: {ROOT}")
add_paragraph(
    document,
    "Для публикации в удаленном репозитории необходимо выполнить инициализацию Git, создать ветки "
    "manual-mvvm и communitytoolkit-mvvm, затем вставить URL репозитория в этот раздел.",
)

document.add_heading("Заключение", level=1)
add_paragraph(
    document,
    "В ходе лабораторной работы разработано WPF-приложение с пятью вкладками. Содержимое вкладок "
    "вынесено в отдельные UserControl. Реализована визуальная модель без сторонних библиотек, "
    "кнопки работают через ICommand, продемонстрированы основные режимы Binding, UpdateSourceTrigger "
    "и варианты триггеров WPF.",
)

document.save(REPORT_PATH)
print(REPORT_PATH)
